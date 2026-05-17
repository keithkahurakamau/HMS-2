"""Generate MediFleet Business Proposal as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0D, 0x2B, 0x55)   # dark navy
TEAL      = RGBColor(0x00, 0x87, 0x8A)   # MediFleet teal
LIGHT_BG  = RGBColor(0xF0, 0xF7, 0xF8)  # very light teal tint
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TEXT = RGBColor(0x44, 0x4D, 0x56)
ACCENT    = RGBColor(0xFF, 0x6B, 0x35)   # warm orange accent


def hex_to_rgb_str(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_to_rgb_str(color))
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            borders.append(el)
    tcPr.append(borders)


def add_run(para, text, bold=False, size=11, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_colored_heading(doc, text, level=1, color=NAVY, size=22, before=200, after=80):
    para = doc.add_paragraph()
    set_para_spacing(para, before=before, after=after)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return para


def add_divider(doc, color=TEAL):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=60)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_to_rgb_str(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

# Default body style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = GREY_TEXT


# ════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER & EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

# ── Cover banner ─────────────────────────────────────────────────────────────
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = banner.cell(0, 0)
set_cell_bg(banner_cell, NAVY)
banner_cell.width = Inches(6.3)

bp = banner_cell.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(bp, before=240, after=40)
add_run(bp, "MediFleet", bold=True, size=36, color=WHITE)
bp2 = banner_cell.add_paragraph()
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(bp2, before=0, after=40)
add_run(bp2, "Hospital Management System", bold=False, size=16, color=RGBColor(0xA8, 0xD8, 0xDA))
bp3 = banner_cell.add_paragraph()
bp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(bp3, before=0, after=200)
add_run(bp3, "Business Proposal  ·  May 2026", italic=True, size=11, color=RGBColor(0xCC, 0xE5, 0xE6))

doc.add_paragraph()  # spacer

# ── Tag line ─────────────────────────────────────────────────────────────────
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(tag, before=60, after=160)
add_run(tag, "Powering Modern Healthcare — One Hospital at a Time", italic=True, size=12, color=TEAL)

add_divider(doc)

# ── Section: Executive Summary ────────────────────────────────────────────────
add_colored_heading(doc, "Executive Summary", size=18, color=NAVY, before=120, after=60)

exec_body = (
    "MediFleet is a cloud-native, multi-tenant Hospital Management System (HMS) "
    "purpose-built for East African healthcare providers. Delivered entirely as a "
    "Software-as-a-Service (SaaS) platform, MediFleet eliminates the capital cost "
    "of on-premise servers and the operational burden of manual software upgrades. "
    "Each subscribing hospital receives a fully isolated database environment, "
    "ensuring complete data privacy and regulatory compliance with the Kenya Data "
    "Protection Act 2019 (KDPA) and the Health Act 2017."
)
p = doc.add_paragraph(exec_body)
set_para_spacing(p, before=0, after=120)

exec_body2 = (
    "The platform ships with a configurable module catalogue of 23 à-la-carte "
    "features — from patient registration and clinical encounters to M-Pesa mobile "
    "payments, pharmacy dispensing, laboratory results, radiology reporting, and an "
    "integrated patient self-service portal. Hospitals activate only the modules they "
    "need, keeping costs predictable while retaining the ability to scale instantly "
    "as their services grow."
)
p2 = doc.add_paragraph(exec_body2)
set_para_spacing(p2, before=0, after=120)

exec_body3 = (
    "Built on a modern technology stack (FastAPI · React 19 · PostgreSQL · Redis), "
    "MediFleet delivers real-time dashboards, role-based access control, WebSocket "
    "notifications, and automated regulatory reporting — capabilities that previously "
    "required bespoke, six-figure software projects. We are now entering a structured "
    "commercial growth phase and seeking to expand our installed base from pilot "
    "clients to 50 contracted hospitals within 24 months."
)
p3 = doc.add_paragraph(exec_body3)
set_para_spacing(p3, before=0, after=160)

# Key highlights table
highlights = [
    ("Target Market",   "Private hospitals, mission hospitals, and specialist clinics in Kenya and the broader East Africa region"),
    ("Revenue Model",   "Annual SaaS subscriptions + module add-ons + one-time implementation fees"),
    ("Competitive Edge","KDPA-compliant multi-tenancy, M-Pesa-native billing, and a zero-CapEx deployment model"),
    ("Year-1 Target",   "15 contracted hospitals · KES 22.5 M ARR"),
    ("Year-2 Target",   "50 contracted hospitals · KES 78 M ARR"),
    ("Team",            "Founding CTO (FullStack) + 2 backend engineers + 1 DevOps + 2 clinical implementation consultants"),
]

hl_table = doc.add_table(rows=len(highlights), cols=2)
hl_table.style = "Table Grid"
hl_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(highlights):
    row = hl_table.rows[i]
    # label cell
    lc = row.cells[0]
    lc.width = Inches(1.8)
    set_cell_bg(lc, LIGHT_BG)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(lp, before=60, after=60)
    add_run(lp, label, bold=True, size=10, color=NAVY)
    # value cell
    vc = row.cells[1]
    vc.width = Inches(4.5)
    vp = vc.paragraphs[0]
    set_para_spacing(vp, before=60, after=60)
    add_run(vp, value, size=10, color=GREY_TEXT)

doc.add_paragraph()

# ── Page break → Page 2 ───────────────────────────────────────────────────────
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# PAGE 2 — PRICING TABLE & ADD-ON EXPENSES
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "Subscription Pricing", size=18, color=NAVY, before=40, after=60)
add_divider(doc)

subtitle = doc.add_paragraph()
set_para_spacing(subtitle, before=80, after=140)
add_run(subtitle, "All plans are billed annually in Kenya Shillings (KES). "
        "Prices are per-hospital per-year. Implementation fee is a one-time charge.", italic=True, size=10, color=GREY_TEXT)

# ── 3-tier pricing table ──────────────────────────────────────────────────────
tiers = [
    {
        "name":  "Starter",
        "price": "KES 900,000 / yr",
        "tagline": "Community Clinics & Dispensaries",
        "modules": [
            "Patient Registry (unlimited patients)",
            "Appointments & Scheduling",
            "Clinical Desk (Encounters & SOAP Notes)",
            "Internal Messaging & Notifications",
            "Role-Based Access Control (up to 15 users)",
            "Email & In-App Support",
        ],
        "impl": "KES 50,000",
        "color": TEAL,
        "highlight": False,
    },
    {
        "name":  "Professional",
        "price": "KES 1,800,000 / yr",
        "tagline": "Private Hospitals & Specialist Centres",
        "modules": [
            "Everything in Starter",
            "Laboratory Module (orders, results, FEFO reagents)",
            "Pharmacy & Inventory (dispensing, stock alerts)",
            "Wards & In-Patient Management (bed board)",
            "Billing & M-Pesa STK Push Integration",
            "Medical History & KDPA Consent Tracking",
            "Referrals & Analytics Dashboard",
            "Up to 50 users · Priority Support (4-hr SLA)",
        ],
        "impl": "KES 120,000",
        "color": NAVY,
        "highlight": True,   # featured tier
    },
    {
        "name":  "Enterprise",
        "price": "KES 3,600,000 / yr",
        "tagline": "Hospital Groups & Mission Networks",
        "modules": [
            "Everything in Professional",
            "Radiology Module (DICOM viewer integration)",
            "Patient Self-Service Portal",
            "Multi-Branch / Group Reporting",
            "Custom Branding & Print Templates",
            "Insurance Claim Workflow",
            "Unlimited users · Dedicated CSM",
            "24 / 7 Phone + SLA guarantee (1-hr critical)",
            "Annual KDPA compliance audit report",
        ],
        "impl": "KES 250,000",
        "color": RGBColor(0x1A, 0x3A, 0x6B),
        "highlight": False,
    },
]

pricing_table = doc.add_table(rows=1, cols=3)
pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for col_idx, tier in enumerate(tiers):
    cell = pricing_table.cell(0, col_idx)
    cell.width = Inches(2.1)

    # header block
    hp = cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(hp, before=120, after=20)
    set_cell_bg(cell, tier["color"])
    add_run(hp, tier["name"], bold=True, size=16, color=WHITE)

    # price
    pp = cell.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(pp, before=20, after=10)
    add_run(pp, tier["price"], bold=True, size=13, color=RGBColor(0xFF, 0xE0, 0x7A) if tier["highlight"] else WHITE)

    # tagline
    tp = cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(tp, before=0, after=80)
    add_run(tp, tier["tagline"], italic=True, size=9,
            color=RGBColor(0xCC, 0xE5, 0xE6) if not tier["highlight"] else RGBColor(0xA8, 0xD8, 0xDA))

    # divider line in cell (visual separator)
    divp = cell.add_paragraph("─" * 22)
    divp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(divp, before=0, after=60)
    for run in divp.runs:
        run.font.color.rgb = RGBColor(0x88, 0xBB, 0xBD)
        run.font.size = Pt(8)

    # feature list
    for feat in tier["modules"]:
        fp = cell.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(fp, before=30, after=30)
        add_run(fp, "✓  ", bold=True, size=9,
                color=RGBColor(0xFF, 0xE0, 0x7A) if tier["highlight"] else RGBColor(0x7D, 0xE8, 0xC5))
        add_run(fp, feat, size=9, color=WHITE)

    # implementation fee
    impl_p = cell.add_paragraph()
    impl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(impl_p, before=80, after=60)
    add_run(impl_p, f"Implementation fee: {tier['impl']}", italic=True, size=8,
            color=RGBColor(0xAA, 0xCC, 0xCE))

    # "Most Popular" badge on middle tier
    if tier["highlight"]:
        badge_p = cell.add_paragraph()
        badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(badge_p, before=0, after=100)
        add_run(badge_p, "★  MOST POPULAR  ★", bold=True, size=9, color=RGBColor(0xFF, 0xE0, 0x7A))


doc.add_paragraph()  # spacer

# ── Add-on expenses table ─────────────────────────────────────────────────────
add_colored_heading(doc, "Optional Add-On Expenses", size=15, color=TEAL, before=140, after=60)

addons = [
    ("Additional User Licences",    "KES 12,000",  "per user / year",   "For hospitals exceeding plan user limits"),
    ("SMS Notification Bundle",     "KES 25,000",  "per year (5,000 SMS)", "Appointment reminders & prescription alerts"),
    ("DICOM Storage (Radiology)",   "KES 60,000",  "per TB / year",     "Cloud DICOM image archive (Enterprise only)"),
    ("Insurance Claims Module",     "KES 90,000",  "per year",          "SHA / NHIF claim submission & reconciliation"),
    ("Custom Report Development",   "KES 80,000",  "one-time",          "Bespoke BI report built to hospital specs"),
    ("Extended Data Retention",     "KES 30,000",  "per year",          "Beyond standard 7-year retention tier"),
    ("On-Site Training (per day)",  "KES 35,000",  "per day + travel",  "Live staff training at hospital premises"),
    ("Priority Escalation Retainer","KES 150,000", "per year",          "Guaranteed 30-min response for critical issues"),
]

ao_table = doc.add_table(rows=len(addons) + 1, cols=4)
ao_table.style = "Table Grid"
ao_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
headers = ["Add-On", "Price", "Unit", "Description"]
for i, h in enumerate(headers):
    hc = ao_table.cell(0, i)
    set_cell_bg(hc, NAVY)
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(hp, before=60, after=60)
    add_run(hp, h, bold=True, size=9.5, color=WHITE)

# Data rows
for row_idx, (name, price, unit, desc) in enumerate(addons):
    row = ao_table.rows[row_idx + 1]
    bg = LIGHT_BG if row_idx % 2 == 0 else WHITE

    data = [name, price, unit, desc]
    for col_idx, val in enumerate(data):
        cell = row.cells[col_idx]
        set_cell_bg(cell, bg)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(cp, before=50, after=50)
        bold = col_idx in (0, 1)
        color = NAVY if col_idx == 1 else GREY_TEXT
        add_run(cp, val, bold=bold, size=9, color=color)

doc.add_paragraph()

# ── Page break → Page 3 ───────────────────────────────────────────────────────
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# PAGE 3 — REVENUE PROJECTIONS & CLOSING
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "Revenue Projections & Growth Estimates", size=18, color=NAVY, before=40, after=60)
add_divider(doc)

proj_intro = doc.add_paragraph(
    "The following projections are based on conservative market penetration assumptions: "
    "a 36-month sales cycle for government-linked facilities, 18-month cycle for private "
    "hospitals, and an average contract value (ACV) weighted across the three tiers. "
    "Churn is modelled at 5 % annually after year 1."
)
set_para_spacing(proj_intro, before=80, after=120)

# ── Annual projections table ──────────────────────────────────────────────────
projections = [
    # Year | Starter | Professional | Enterprise | New hospitals | Total ARR | Add-on rev | Total rev
    ("Year 1 (2026–27)", "8",  "5",  "2",  "15", "KES 22.5 M",  "KES 3.2 M",  "KES 25.7 M"),
    ("Year 2 (2027–28)", "18", "22", "10", "35", "KES 57.6 M",  "KES 9.5 M",  "KES 67.1 M"),
    ("Year 3 (2028–29)", "25", "40", "20", "50", "KES 117.0 M", "KES 21.0 M", "KES 138.0 M"),
]
proj_headers = ["Period", "Starter\nClients", "Professional\nClients", "Enterprise\nClients",
                "Total\nHospitals", "Subscription\nARR", "Add-On\nRevenue", "Total\nRevenue"]

proj_table = doc.add_table(rows=len(projections) + 1, cols=len(proj_headers))
proj_table.style = "Table Grid"
proj_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(proj_headers):
    hc = proj_table.cell(0, i)
    set_cell_bg(hc, TEAL)
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(hp, before=60, after=60)
    add_run(hp, h, bold=True, size=8.5, color=WHITE)

row_colors = [LIGHT_BG, WHITE, RGBColor(0xE8, 0xF5, 0xF5)]
for r_idx, row_data in enumerate(projections):
    row = proj_table.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        set_cell_bg(cell, row_colors[r_idx])
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(cp, before=55, after=55)
        bold = c_idx in (0, 7)
        color = NAVY if c_idx == 7 else GREY_TEXT
        add_run(cp, val, bold=bold, size=9, color=color)

doc.add_paragraph()

# ── Cost structure & margin ───────────────────────────────────────────────────
add_colored_heading(doc, "Estimated Operating Cost Structure", size=14, color=TEAL, before=120, after=60)

costs = [
    ("Cloud Infrastructure (AWS/Render)",    "KES 3.6 M",  "KES 7.2 M",  "KES 12.0 M"),
    ("Engineering Salaries (5 FTE → 12 FTE)","KES 9.0 M",  "KES 18.0 M", "KES 28.0 M"),
    ("Sales & Marketing",                    "KES 2.5 M",  "KES 8.0 M",  "KES 14.0 M"),
    ("Implementation & Customer Success",    "KES 1.8 M",  "KES 5.0 M",  "KES 9.0 M"),
    ("Legal, Compliance & KDPA Audit",       "KES 0.8 M",  "KES 1.2 M",  "KES 1.5 M"),
    ("G&A (Finance, HR, Office)",            "KES 1.2 M",  "KES 2.0 M",  "KES 3.0 M"),
    ("TOTAL COSTS",                          "KES 18.9 M", "KES 41.4 M", "KES 67.5 M"),
    ("GROSS PROFIT (est.)",                  "KES 6.8 M",  "KES 25.7 M", "KES 70.5 M"),
    ("GROSS MARGIN",                         "26 %",       "38 %",       "51 %"),
]

cost_table = doc.add_table(rows=len(costs) + 1, cols=4)
cost_table.style = "Table Grid"
cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cost_headers = ["Cost Category", "Year 1", "Year 2", "Year 3"]
for i, h in enumerate(cost_headers):
    hc = cost_table.cell(0, i)
    set_cell_bg(hc, NAVY)
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(hp, before=55, after=55)
    add_run(hp, h, bold=True, size=9, color=WHITE)

for r_idx, row_data in enumerate(costs):
    row = cost_table.rows[r_idx + 1]
    is_total_row = row_data[0].startswith("TOTAL") or row_data[0].startswith("GROSS")
    bg = NAVY if is_total_row else (LIGHT_BG if r_idx % 2 == 0 else WHITE)
    txt_color = WHITE if is_total_row else GREY_TEXT
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        set_cell_bg(cell, bg)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(cp, before=50, after=50)
        add_run(cp, val, bold=is_total_row, size=9, color=txt_color)

doc.add_paragraph()

# ── Market opportunity note ───────────────────────────────────────────────────
add_colored_heading(doc, "Market Opportunity", size=14, color=TEAL, before=100, after=60)

market_para = doc.add_paragraph(
    "Kenya has approximately 9,800 registered health facilities (Ministry of Health 2023), "
    "of which roughly 2,400 are private hospitals and specialist clinics — our primary "
    "addressable segment. Penetrating just 2 % of this market (≈ 48 facilities) achieves "
    "Year-3 targets. The broader East Africa region (Uganda, Tanzania, Rwanda) adds an "
    "estimated 4,200 comparable facilities, representing a medium-term expansion opportunity "
    "exceeding KES 1 B in addressable ARR."
)
set_para_spacing(market_para, before=0, after=100)

# ── Closing call to action ─────────────────────────────────────────────────────
cta_table = doc.add_table(rows=1, cols=1)
cta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cta_cell = cta_table.cell(0, 0)
set_cell_bg(cta_cell, TEAL)

cp1 = cta_cell.paragraphs[0]
cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(cp1, before=120, after=30)
add_run(cp1, "Ready to modernise your hospital?", bold=True, size=13, color=WHITE)

cp2 = cta_cell.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(cp2, before=0, after=30)
add_run(cp2, "Contact our team for a free 30-day pilot and live demo.", size=10.5, color=RGBColor(0xCC, 0xF0, 0xF2))

cp3 = cta_cell.add_paragraph()
cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(cp3, before=0, after=100)
add_run(cp3, "sales@medifleet.co.ke  ·  +254 700 000 000  ·  www.medifleet.co.ke",
        bold=True, size=10, color=WHITE)

doc.add_paragraph()

# ── Footer disclaimer ──────────────────────────────────────────────────────────
disc = doc.add_paragraph(
    "This proposal is confidential and intended solely for the addressee. Revenue projections are estimates based "
    "on current market data and internal assumptions; actual results may vary. All prices exclusive of VAT. "
    "Prices valid for 90 days from date of issue.  |  MediFleet © 2026"
)
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(disc, before=80, after=0)
for run in disc.runs:
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x99, 0xAA, 0xBB)
    run.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/home/user/HMS-2/MediFleet_Business_Proposal_2026.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
